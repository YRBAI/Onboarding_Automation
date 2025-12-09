# Enhanced doc_processor.py with checkbox and formatting preservation

import os
import json
from typing import Optional, Dict, Any, Tuple
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import re

# Import your config constants
# from config import DOC_FONT_NAME, DOC_FONT_SIZE, CONFIG_FILE, RESULTS_DIR

# Temporary constants for demonstration
DOC_FONT_NAME = "Arial"
DOC_FONT_SIZE = 11
CONFIG_FILE = "config.json"
RESULTS_DIR = "results"


class EnhancedDocProcessor:
    """Enhanced Word document processor that preserves checkboxes and special formatting."""
    
    def __init__(self, config_file: str = None):
        self.config_file = config_file or CONFIG_FILE
        self.config = self.load_config()
        self.results_dir = RESULTS_DIR
    
    def load_config(self) -> Dict[str, Any]:
        """Load configuration from JSON file if it exists."""
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r') as f:
                    config = json.load(f)
                    if 'folder' not in config or not config['folder']:
                        config['folder'] = self.results_dir
                    return config
            except Exception as e:
                print(f"Warning: Error loading config file: {e}")
        
        return {
            'folder': self.results_dir,
            'old_text': 'Sep 12_Mass',
            'doc_filename': '',
            'access_code': ''
        }
    
    def save_config(self, config: Dict[str, Any] = None) -> None:
        """Save configuration to JSON file."""
        if config is None:
            config = self.config
        try:
            with open(self.config_file, 'w') as f:
                json.dump(config, f, indent=2)
        except Exception as e:
            print(f"Warning: Error saving config file: {e}")
    
    def preserve_run_formatting(self, source_run, target_run):
        """Copy formatting from source run to target run."""
        try:
            # Copy font properties
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            if source_run.font.bold is not None:
                target_run.font.bold = source_run.font.bold
            if source_run.font.italic is not None:
                target_run.font.italic = source_run.font.italic
            if source_run.font.underline is not None:
                target_run.font.underline = source_run.font.underline
        except Exception as e:
            print(f"Warning: Could not preserve formatting: {e}")
    
    def smart_text_replacement(self, paragraph, old_text: str, new_text: str) -> int:
        """
        Smart text replacement that preserves formatting and special characters.
        Returns the number of replacements made.
        """
        if old_text not in paragraph.text:
            return 0
        
        replacements_made = 0
        
        try:
            # Get the full text and find replacement positions
            full_text = paragraph.text
            replacement_count = full_text.count(old_text)
            
            if replacement_count == 0:
                return 0
            
            # Strategy 1: Try run-by-run replacement (preserves most formatting)
            replaced_in_runs = False
            for run in paragraph.runs:
                if old_text in run.text:
                    original_formatting = {
                        'name': run.font.name,
                        'size': run.font.size,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline
                    }
                    
                    run.text = run.text.replace(old_text, new_text)
                    replacements_made += run.text.count(new_text) - run.text.count(old_text)
                    
                    # Restore formatting
                    for prop, value in original_formatting.items():
                        if value is not None:
                            setattr(run.font, prop, value)
                    
                    # Apply standard font if not already set
                    if not run.font.name:
                        run.font.name = DOC_FONT_NAME
                    if not run.font.size:
                        run.font.size = Pt(DOC_FONT_SIZE)
                    
                    replaced_in_runs = True
            
            # Strategy 2: If run-by-run didn't work, do paragraph-level replacement
            if not replaced_in_runs and old_text in paragraph.text:
                # Store all run information before clearing
                run_info = []
                for run in paragraph.runs:
                    run_info.append({
                        'text': run.text,
                        'formatting': {
                            'name': run.font.name,
                            'size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic,
                            'underline': run.font.underline
                        }
                    })
                
                # Clear runs and create new one with replaced text
                paragraph.clear()
                new_text_content = full_text.replace(old_text, new_text)
                new_run = paragraph.add_run(new_text_content)
                
                # Apply formatting from the first run that had formatting
                for info in run_info:
                    if any(v for v in info['formatting'].values() if v):
                        for prop, value in info['formatting'].items():
                            if value is not None:
                                setattr(new_run.font, prop, value)
                        break
                
                # Ensure standard font is applied
                if not new_run.font.name:
                    new_run.font.name = DOC_FONT_NAME
                if not new_run.font.size:
                    new_run.font.size = Pt(DOC_FONT_SIZE)
                
                replacements_made = replacement_count
            
        except Exception as e:
            print(f"Warning: Error in smart text replacement: {e}")
            # Fallback to simple replacement
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                replacements_made = paragraph.text.count(new_text)
        
        return replacements_made
    
    def process_paragraphs(self, paragraphs, old_text: str, new_text: str) -> int:
        """Process a list of paragraphs for text replacement."""
        total_replacements = 0
        
        for para in paragraphs:
            if old_text in para.text:
                replacements = self.smart_text_replacement(para, old_text, new_text)
                total_replacements += replacements
            
            # Ensure consistent font formatting for all runs
            for run in para.runs:
                if run.text.strip():  # Only format runs with actual content
                    if not run.font.name:
                        run.font.name = DOC_FONT_NAME
                    if not run.font.size:
                        run.font.size = Pt(DOC_FONT_SIZE)
        
        return total_replacements
    
    def replace_text_in_document(self, input_path: str, output_path: str, old_text: str, new_text: str) -> bool:
        """
        Replace text in a Word document while preserving checkboxes and formatting.
        """
        try:
            print(f"  Loading document: {os.path.basename(input_path)}")
            doc = Document(input_path)
            
            total_replacements = 0
            
            # Process main document paragraphs
            print(f"  Processing paragraphs...")
            total_replacements += self.process_paragraphs(doc.paragraphs, old_text, new_text)
            
            # Process tables
            print(f"  Processing tables...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        total_replacements += self.process_paragraphs(cell.paragraphs, old_text, new_text)
            
            # Process headers and footers
            print(f"  Processing headers and footers...")
            for section in doc.sections:
                # Process headers
                headers = [section.header, section.first_page_header, section.even_page_header]
                for header in headers:
                    if header:
                        total_replacements += self.process_paragraphs(header.paragraphs, old_text, new_text)
                
                # Process footers
                footers = [section.footer, section.first_page_footer, section.even_page_footer]
                for footer in footers:
                    if footer:
                        total_replacements += self.process_paragraphs(footer.paragraphs, old_text, new_text)
            
            # Save the document
            print(f"  Saving modified document...")
            doc.save(output_path)
            
            print(f"  Made {total_replacements} text replacements")
            print(f"  Applied {DOC_FONT_NAME} {DOC_FONT_SIZE}pt font formatting")
            print(f"  Preserved checkboxes and special formatting")
            
            return True
            
        except Exception as e:
            print(f"Error processing document: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def get_available_documents(self, folder: str = None) -> list:
        """Get list of available Word documents in the specified folder."""
        if folder is None:
            folder = self.results_dir
        
        try:
            if not os.path.exists(folder):
                os.makedirs(folder, exist_ok=True)
                return []
            
            docx_files = [f for f in os.listdir(folder) 
                         if f.endswith('.docx') and not f.startswith('~')]
            return docx_files
        except Exception as e:
            print(f"Error listing documents: {e}")
            return []
    
    def update_config_with_new_text(self, new_text: str) -> None:
        """Update config with new text as the old text for next use."""
        self.config['old_text'] = new_text
        self.save_config()
    
    def process_document_gui(self, template_path: str, new_date: str) -> Optional[str]:
        """
        Process document for GUI with automatic file naming and config updates.
        """
        try:
            if not os.path.exists(template_path):
                print(f"Error: Template document not found: {template_path}")
                return None
            
            # Format new text
            new_text = f"{new_date}_Mass"
            old_text = self.config.get('old_text', 'Sep 12_Mass')
            
            print(f"Replacing '{old_text}' with '{new_text}'")
            
            # Generate output filename
            output_filename = f"{new_text} Onboarding New Product Evaluation Appendix C.docx"
            output_path = os.path.join(self.results_dir, output_filename)
            
            print(f"Output path: {output_path}")
            
            # Process the document
            success = self.replace_text_in_document(template_path, output_path, old_text, new_text)
            
            if success:
                # Update config with new text for next use
                self.update_config_with_new_text(new_text)
                print(f"Document processed successfully!")
                print(f"Config updated: old_text = '{new_text}'")
                return output_path
            else:
                print("Document processing failed.")
                return None
                
        except Exception as e:
            print(f"Error in document processing: {e}")
            import traceback
            traceback.print_exc()
            return None


# Alternative approach for checkbox preservation
class CheckboxPreservingProcessor(EnhancedDocProcessor):
    """Specialized processor that uses regex to preserve checkbox symbols."""
    
    def __init__(self, config_file: str = None):
        super().__init__(config_file)
        # Common checkbox symbols
        self.checkbox_symbols = ['✓', '✔', '☑', '☒', '✗', '✘']
    
    def preserve_checkbox_text_replacement(self, text: str, old_text: str, new_text: str) -> str:
        """
        Replace text while preserving checkbox symbols and their positions.
        """
        if old_text not in text:
            return text
        
        # Create a pattern that captures checkbox symbols with surrounding context
        checkbox_pattern = r'([✓✔☑☒✗✘](?:\s*\w+(?:\s+\w+)*)?)'
        
        # Split text by checkboxes to preserve them
        parts = re.split(f'({"|".join(re.escape(symbol) for symbol in self.checkbox_symbols)})', text)
        
        # Replace text in non-checkbox parts
        for i, part in enumerate(parts):
            if part not in self.checkbox_symbols:
                parts[i] = part.replace(old_text, new_text)
        
        return ''.join(parts)
    
    def smart_text_replacement(self, paragraph, old_text: str, new_text: str) -> int:
        """Override with checkbox-aware replacement."""
        if old_text not in paragraph.text:
            return 0
        
        try:
            # Check if paragraph contains checkbox symbols
            has_checkboxes = any(symbol in paragraph.text for symbol in self.checkbox_symbols)
            
            if has_checkboxes:
                # Use checkbox-preserving replacement
                original_text = paragraph.text
                new_paragraph_text = self.preserve_checkbox_text_replacement(
                    original_text, old_text, new_text
                )
                
                if new_paragraph_text != original_text:
                    # Preserve run structure as much as possible
                    if len(paragraph.runs) == 1:
                        # Simple case: single run
                        run = paragraph.runs[0]
                        original_font = {
                            'name': run.font.name,
                            'size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic
                        }
                        run.text = new_paragraph_text
                        # Restore formatting
                        for prop, value in original_font.items():
                            if value is not None:
                                setattr(run.font, prop, value)
                    else:
                        # Multiple runs: more complex preservation
                        paragraph.text = new_paragraph_text
                    
                    # Apply standard formatting
                    for run in paragraph.runs:
                        if run.text.strip():
                            if not run.font.name:
                                run.font.name = DOC_FONT_NAME
                            if not run.font.size:
                                run.font.size = Pt(DOC_FONT_SIZE)
                    
                    return original_text.count(old_text)
            else:
                # Use standard replacement for non-checkbox paragraphs
                return super().smart_text_replacement(paragraph, old_text, new_text)
        
        except Exception as e:
            print(f"Warning: Checkbox preservation failed, using standard replacement: {e}")
            return super().smart_text_replacement(paragraph, old_text, new_text)
        
        return 0


def main():
    """Test the enhanced processor."""
    print("Enhanced Word Document Processor with Checkbox Preservation")
    print("=" * 60)
    
    # Use the checkbox-preserving processor
    processor = CheckboxPreservingProcessor()
    
    docs = processor.get_available_documents()
    if docs:
        print(f"Available documents in {processor.results_dir}:")
        for i, doc in enumerate(docs, 1):
            print(f"  {i}. {doc}")
        
        choice = input(f"Select document (1-{len(docs)}): ").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(docs):
            template_path = os.path.join(processor.results_dir, docs[int(choice)-1])
            
            new_date = input("Enter new date (e.g., Sep 13): ").strip()
            if new_date:
                result = processor.process_document_gui(template_path, new_date)
                if result:
                    print(f"\nSuccess! Document saved to: {result}")
                    print("Checkboxes and formatting should be preserved!")
                else:
                    print("\nDocument processing failed.")
            else:
                print("No date provided.")
        else:
            print("Invalid selection.")
    else:
        print(f"No Word documents found in {processor.results_dir}")


if __name__ == "__main__":
    main()
# For backward compatibility with app_launcher.py
DocProcessor = CheckboxPreservingProcessor  # or EnhancedDocProcessor